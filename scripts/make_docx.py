"""Render report/final_report.md to a formatted Word document.

Hand-rolled rather than pandoc-based: pandoc/LibreOffice are not installable on this
machine (no sudo), so this walks the subset of Markdown the report actually uses --
headings, paragraphs, tables, bullet/numbered lists, blockquotes, rules and inline
bold/italic/code -- and emits styled python-docx output.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "report" / "final_report.md"
OUT = ROOT / "report" / "final_report.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
MIDBLUE = RGBColor(0x2C, 0x5F, 0x8D)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", re.S)


def shade(cell, hexcolor):
    """Apply a solid background fill to a table cell."""
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def add_runs(par, text):
    """Write text into a paragraph, honouring **bold**, *italic* and `code`."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        else:
            par.add_run(part)


def is_table_sep(line):
    return bool(re.match(r"^\|[\s:|-]+\|$", line.strip())) and "-" in line


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(re.sub(r"\*\*", "", h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, "1F4E79")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row[: len(header)]):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_runs(p, val)
            for r in p.runs:
                r.font.size = Pt(9)
            if ri % 2 == 1:
                shade(cells[i], "F5F8FB")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def build():
    lines = SRC.read_text(encoding="utf-8").split("\n")
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = DARK
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.9)
        sec.left_margin = sec.right_margin = Inches(0.9)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", s):
            p = doc.add_paragraph()
            pr = p._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "D5DDE5")
            bd.append(bot)
            pr.append(bd)
            i += 1
            continue

        # tables
        if s.startswith("|") and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(s)
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            level, text = len(m.group(1)), m.group(2)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            if level == 1:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(19)
                r.font.color.rgb = NAVY
                p.paragraph_format.space_after = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(14 if level == 2 else 11.5)
                r.font.color.rgb = NAVY if level == 2 else MIDBLUE
                p.paragraph_format.space_before = Pt(14 if level == 2 else 10)
                p.paragraph_format.space_after = Pt(5)
                p.paragraph_format.keep_with_next = True
                if level == 2:
                    pr = p._p.get_or_add_pPr()
                    bd = OxmlElement("w:pBdr")
                    bot = OxmlElement("w:bottom")
                    bot.set(qn("w:val"), "single")
                    bot.set(qn("w:sz"), "4")
                    bot.set(qn("w:color"), "C8D6E5")
                    bd.append(bot)
                    pr.append(bd)
            i += 1
            continue

        # blockquote
        if s.startswith(">"):
            text = re.sub(r"^>\s?", "", s)
            text = re.sub(r"^\*(.+)\*$", r"\1", text.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            add_runs(p, text)
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x5A, 0x4A, 0x2A)
            i += 1
            continue

        # bullets / numbered
        mb = re.match(r"^[-*]\s+(.*)$", s)
        mn = re.match(r"^(\d+)\.\s+(.*)$", s)
        if mb or mn:
            style = "List Bullet" if mb else "List Number"
            p = doc.add_paragraph(style=style)
            add_runs(p, (mb or mn).group(1 if mb else 2))
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # paragraph (join wrapped lines)
        buf = [s]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,4}\s|\||>|[-*]\s|\d+\.\s|-{3,}$)", lines[i].strip()
        ):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, " ".join(buf))

    doc.save(OUT)
    print(f"wrote {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    build()
